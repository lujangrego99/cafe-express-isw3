#!/usr/bin/env python3
"""Genera el DOCX del informe a partir del mismo Markdown.

Parser acotado al subconjunto de Markdown usado en el informe:
encabezados, tablas, imagenes, listas, bloques de codigo, citas e inline
(**negrita**, `codigo`, [texto](url)).

Uso:  python3 build_docx.py
Salida: TP-Final-Integrador_CafeExpress.docx
"""
import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from PIL import Image

AQUI = os.path.dirname(os.path.abspath(__file__))
MD = os.path.join(AQUI, "TP-Final-Integrador_CafeExpress.md")
DOCX = os.path.join(AQUI, "TP-Final-Integrador_CafeExpress.docx")
MAX_IN = 6.3  # ancho util de pagina A4 con margenes

INLINE = re.compile(r"(\*\*.+?\*\*|`.+?`|\[.+?\]\(.+?\))")
IMG = re.compile(r"^!\[(.*?)\]\((.+?)\)\s*$")


def add_runs(parrafo, texto):
    """Agrega texto con formato inline (negrita, codigo, links)."""
    for token in INLINE.split(texto):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**"):
            parrafo.add_run(token[2:-2]).bold = True
        elif token.startswith("`") and token.endswith("`"):
            r = parrafo.add_run(token[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(9.5)
        elif token.startswith("["):
            m = re.match(r"\[(.+?)\]\((.+?)\)", token)
            if m:
                r = parrafo.add_run(m.group(1) + " (" + m.group(2) + ")")
                r.font.color.rgb = RGBColor(0x1F, 0x5E, 0x8C)
        else:
            parrafo.add_run(token)


def add_image(doc, ruta):
    full = os.path.normpath(os.path.join(AQUI, ruta))
    if not os.path.exists(full):
        doc.add_paragraph(f"[imagen no encontrada: {ruta}]")
        return
    try:
        with Image.open(full) as im:
            ancho_in = im.width / 96.0
    except Exception:
        ancho_in = MAX_IN
    doc.add_picture(full, width=Inches(min(MAX_IN, ancho_in)))
    doc.paragraphs[-1].alignment = 1  # centrado


def parse_tabla(lineas, i, doc):
    """Procesa una tabla Markdown desde la linea i. Devuelve el nuevo indice."""
    filas = []
    while i < len(lineas) and lineas[i].strip().startswith("|"):
        filas.append([c.strip() for c in lineas[i].strip().strip("|").split("|")])
        i += 1
    # filas[1] es el separador ---
    encabezado = filas[0]
    cuerpo = filas[2:]
    tabla = doc.add_table(rows=1, cols=len(encabezado))
    tabla.style = "Light Grid Accent 1"
    for j, celda in enumerate(encabezado):
        p = tabla.rows[0].cells[j].paragraphs[0]
        add_runs(p, celda)
        for run in p.runs:
            run.bold = True
    for fila in cuerpo:
        cells = tabla.add_row().cells
        for j, celda in enumerate(fila):
            if j < len(cells):
                add_runs(cells[j].paragraphs[0], celda)
    doc.add_paragraph()
    return i


def main():
    with open(MD, encoding="utf-8") as f:
        lineas = f.read().split("\n")

    doc = Document()
    doc.styles["Normal"].font.name = "Segoe UI"
    doc.styles["Normal"].font.size = Pt(11)

    i = 0
    while i < len(lineas):
        linea = lineas[i]
        s = linea.strip()

        if not s:
            i += 1
            continue

        # Bloque de codigo
        if s.startswith("```"):
            i += 1
            codigo = []
            while i < len(lineas) and not lineas[i].strip().startswith("```"):
                codigo.append(lineas[i])
                i += 1
            i += 1  # cierre
            p = doc.add_paragraph()
            r = p.add_run("\n".join(codigo))
            r.font.name = "Consolas"
            r.font.size = Pt(9)
            continue

        # Imagen
        m = IMG.match(s)
        if m:
            add_image(doc, m.group(2))
            i += 1
            continue

        # Tabla
        if s.startswith("|"):
            i = parse_tabla(lineas, i, doc)
            continue

        # Regla horizontal
        if s == "---":
            i += 1
            continue

        # Encabezados
        if s.startswith("#"):
            nivel = len(s) - len(s.lstrip("#"))
            texto = s[nivel:].strip()
            if nivel == 1:
                t = doc.add_heading("", level=0)
                add_runs(t, texto)
            else:
                t = doc.add_heading("", level=min(nivel - 1, 4))
                add_runs(t, texto)
            i += 1
            continue

        # Cita
        if s.startswith(">"):
            p = doc.add_paragraph(style="Intense Quote")
            add_runs(p, s[1:].strip())
            i += 1
            continue

        # Listas (con un nivel de anidacion por sangria)
        m_ul = re.match(r"^(\s*)[-*]\s+(.*)$", linea)
        m_ol = re.match(r"^(\s*)\d+\.\s+(.*)$", linea)
        if m_ul or m_ol:
            m = m_ul or m_ol
            sangria = len(m.group(1))
            anidado = sangria >= 2
            if m_ul:
                estilo = "List Bullet 2" if anidado else "List Bullet"
            else:
                estilo = "List Number 2" if anidado else "List Number"
            p = doc.add_paragraph(style=estilo)
            add_runs(p, m.group(2))
            i += 1
            continue

        # Parrafo normal
        p = doc.add_paragraph()
        add_runs(p, s)
        i += 1

    doc.save(DOCX)
    print("DOCX generado:", DOCX)


if __name__ == "__main__":
    main()
